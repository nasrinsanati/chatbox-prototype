# file_parser.py - Clean Stable Version
import pdfplumber
from pypdf import PdfReader
from docx import Document
import streamlit as st
import re

def clean_extracted_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    lines = [line.strip() for line in text.split('\n') if len(line.strip()) > 2]
    return '\n'.join(lines).strip()


def extract_text_from_pdf(uploaded_file):
    try:
        text = ""
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        return clean_extracted_text(text)
    except Exception as e:
        st.warning(f"pdfplumber failed, falling back to pypdf: {e}")
        try:
            reader = PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return clean_extracted_text(text)
        except Exception as e2:
            st.error(f"PDF extraction failed: {e2}")
            return ""


def extract_text_from_docx(uploaded_file):
    """Extract paragraphs AND tables from a Word syllabus."""
    try:
        doc = Document(uploaded_file)
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        for t_i, table in enumerate(doc.tables, start=1):
            rows = []
            for row in table.rows:
                cells = [" ".join(cell.text.split()) for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                parts.append(f"\n=== TABLE {t_i} ===")
                parts.extend(rows)
                parts.append("=== END TABLE ===\n")

        return clean_extracted_text("\n".join(parts))
    except Exception as e:
        st.error(f"DOCX extraction failed: {e}")
        return ""